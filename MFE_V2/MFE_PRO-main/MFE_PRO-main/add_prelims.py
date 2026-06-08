"""
Adds Dédicaces and Remerciements pages, then updates the Table des matières.
Run against MFE Fodhil & Chibani_Fixed.docx → produces _v2.docx
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.shared import Pt

INPUT  = "MFE Fodhil & Chibani_Fixed.docx"
OUTPUT = "MFE Fodhil & Chibani_v2.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
def w(t): return "{%s}%s" % (W_NS, t)

def para_text(p):
    return "".join(r.text for r in p.runs if r.text)

def ins(ref, text="", bold=False, italic=False, center=False, pb=False, font_size=None):
    """Insert a new paragraph immediately before ref."""
    new_p = OxmlElement("w:p")
    ref._p.addprevious(new_p)
    para = DocxParagraph(new_p, ref._p.getparent())
    if pb:
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(w("type"), "page")
        run._r.append(br)
    elif text is not None:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if font_size:
            run.font.size = Pt(font_size)
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def find_para(doc, needle):
    """Return the first paragraph whose text contains needle."""
    for p in doc.paragraphs:
        if needle in para_text(p):
            return p
    return None

def find_sect_after(doc, after_needle):
    """Return the first paragraph with sectPr that comes after `after_needle` paragraph."""
    found = False
    for p in doc.paragraphs:
        if after_needle in para_text(p):
            found = True
            continue
        if found:
            ppr = p._p.find(w("pPr"))
            if ppr is not None and ppr.find(w("sectPr")) is not None:
                return p
    return None


# ── 1. DÉDICACES ──────────────────────────────────────────────────────────────
def add_dedicaces(toc_title_para):
    print("\n[1] Adding Dédicaces page")

    # Page break — starts the Dédicaces page (pushes off cover)
    ins(toc_title_para, pb=True)

    # Title
    ins(toc_title_para, "")
    ins(toc_title_para, "DÉDICACES", bold=True, center=True, font_size=14)
    ins(toc_title_para, "")
    ins(toc_title_para, "─" * 40, center=True)
    ins(toc_title_para, "")

    # ── Fodhil's dedication ──────────────────────────────────────────────
    ins(toc_title_para,
        "À mes parents, qui ont tout sacrifié pour me permettre d'atteindre ce "
        "niveau. Votre amour et votre soutien indéfectible m'ont donné la force "
        "de persévérer.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À mes frères et sœurs, pour leur présence et leurs encouragements au "
        "quotidien.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À mon binôme et ami Abderrahmane, pour sa collaboration sincère, "
        "sa rigueur et son engagement tout au long de ce projet.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À tous mes amis et camarades de promotion, avec qui j'ai partagé "
        "cette belle aventure académique.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para, "Yacine FODHIL", bold=True,
        center=False)  # right-aligned feel via spacing
    ins(toc_title_para, "")
    ins(toc_title_para, "─" * 40, center=True)
    ins(toc_title_para, "")

    # ── Chibani's dedication ─────────────────────────────────────────────
    ins(toc_title_para,
        "À mes parents, pilier de ma vie, dont le dévouement et les sacrifices "
        "m'ont ouvert la voie de la réussite.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À ma famille, pour l'amour inconditionnel et le soutien moral "
        "qu'elle m'a toujours offerts.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À mon binôme et ami Yacine, pour sa persévérance, son sérieux et "
        "la solidarité dont il a fait preuve durant ce projet.", italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para,
        "À tous ceux qui ont cru en moi et m'ont soutenu de près ou de loin.",
        italic=True)
    ins(toc_title_para, "")
    ins(toc_title_para, "Abderrahmane CHIBANI", bold=True)

    print("  + Dédicaces page inserted")


# ── 2. REMERCIEMENTS ──────────────────────────────────────────────────────────
def add_remerciements(toc_title_para):
    print("\n[2] Adding Remerciements page")

    # Page break — new page for Remerciements
    ins(toc_title_para, pb=True)

    # Title
    ins(toc_title_para, "")
    ins(toc_title_para, "REMERCIEMENTS", bold=True, center=True, font_size=14)
    ins(toc_title_para, "")

    # Paragraphs
    ins(toc_title_para,
        "Nous remercions en premier lieu Allah, le Tout-Puissant, de nous avoir "
        "accordé la volonté, la patience et la santé nécessaires à "
        "l'accomplissement de ce travail.")
    ins(toc_title_para, "")

    ins(toc_title_para,
        "Nous adressons nos plus sincères remerciements à notre encadrante, "
        "Mme Nadia BENAHMED-EL ALIA, et à notre co-encadrant, M. Amar BALLA, "
        "pour leur disponibilité, leurs orientations éclairées, leurs précieux "
        "conseils et l'intérêt constant qu'ils ont porté à notre travail tout "
        "au long de la réalisation de ce projet.")
    ins(toc_title_para, "")

    ins(toc_title_para,
        "Nous tenons également à exprimer notre gratitude à l'ensemble du corps "
        "enseignant de l'Institut de Formation d'Assurances et de Gestion (IFAG) "
        "pour la qualité de la formation dispensée et le cadre académique rigoureux "
        "dont nous avons bénéficié au cours de notre cursus.")
    ins(toc_title_para, "")

    ins(toc_title_para,
        "Nos remerciements vont également à nos familles respectives pour leur "
        "soutien indéfectible, leurs encouragements constants et les sacrifices "
        "consentis afin de nous permettre de mener à bien ce projet.")
    ins(toc_title_para, "")

    ins(toc_title_para,
        "Enfin, nous remercions toutes les personnes qui ont contribué, de près "
        "ou de loin, à la concrétisation de ce mémoire.")

    # Page break before TOC
    ins(toc_title_para, pb=True)

    print("  + Remerciements page inserted")


# ── 3. TABLE DES MATIÈRES — update entries ────────────────────────────────────
def update_toc(toc_sect_para):
    """Insert TOC entries before the sectPr paragraph (just after TOC title)."""
    print("\n[3] Updating Table des matières")

    def toc_entry(text, page, level=0, bold=False):
        indent = "    " * level
        dots_len = max(2, 60 - len(indent) - len(text) - len(str(page)))
        dots = "." * dots_len
        line = "%s%s %s %s" % (indent, text, dots, page)
        ins(toc_sect_para, line, bold=bold)

    ins(toc_sect_para, "")

    # Preliminary sections (Roman numerals)
    toc_entry("Dédicaces",                                      "i",   bold=False)
    toc_entry("Remerciements",                                  "ii",  bold=False)
    toc_entry("Liste des Figures",                              "iii", bold=False)
    toc_entry("Liste des Tableaux",                             "iv",  bold=False)
    toc_entry("Liste des Abréviations",                         "v",   bold=False)
    toc_entry("Note sur l'Utilisation de l'IA",                "vi",  bold=False)
    ins(toc_sect_para, "")

    # Chapter 1
    toc_entry("CHAPITRE 1 : ISO 15118 Standard – Definition & Application",
              "1", bold=True)
    toc_entry("What is ISO 15118?",                            "1",  level=1)
    toc_entry("Why ISO 15118 Matters for Smart EV Charging",   "1",  level=1)
    toc_entry("Application in Our EV Charging Platform",       "2",  level=1)
    toc_entry("System Architecture",                           "2",  level=1)
    toc_entry("DC Charging Mode and CC-CV Strategy",           "3",  level=1)
    toc_entry("Dynamic Parameter Exchange",                    "3",  level=1)
    toc_entry("Demand Response with Photovoltaic Integration", "3",  level=1)
    toc_entry("IoT Integration and Monitoring",                "4",  level=1)
    toc_entry("Security",                                      "4",  level=1)
    toc_entry("Summary",                                       "4",  level=1)
    toc_entry("References",                                    "5",  level=1)
    ins(toc_sect_para, "")

    # Chapter 2
    toc_entry("CHAPITRE 2 : OCPP 1.6 vs. OCPP 2.0 – A Comprehensive Comparison",
              "7", bold=True)
    toc_entry("Device Management and Structure",               "7",  level=1)
    toc_entry("Smart Charging Capabilities",                   "7",  level=1)
    toc_entry("Message Handling and Communication",            "8",  level=1)
    toc_entry("Enhanced Security",                             "9",  level=1)
    toc_entry("Firmware Management",                           "10", level=1)
    toc_entry("Backward Compatibility",                        "10", level=1)
    toc_entry("Use Cases and Scalability",                     "10", level=1)
    toc_entry("Technical Comparison",                          "11", level=1)
    toc_entry("Why Using OCPP 1.6 in Algeria",                "11", level=1)
    toc_entry("Conclusion",                                    "12", level=1)
    toc_entry("References",                                    "12", level=1)
    ins(toc_sect_para, "")

    print("  + TOC entries inserted (page numbers are current estimates)")
    print("  ! Reminder: update page numbers once document is complete")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    shutil.copy2(INPUT, OUTPUT)  # work on a copy
    doc = Document(OUTPUT)

    # ── Anchors ───────────────────────────────────────────────────────────
    toc_title_para = find_para(doc, "Table des matières")
    if not toc_title_para:
        print("ERROR: 'Table des matières' not found"); return
    idx1 = next((i for i,p in enumerate(doc.paragraphs) if p._p is toc_title_para._p), -1)
    print("Anchor: 'Table des matieres' found at para index", idx1)

    toc_sect_para = find_sect_after(doc, "Table des")
    if not toc_sect_para:
        print("ERROR: sectPr after TOC not found"); return
    idx2 = next((i for i,p in enumerate(doc.paragraphs) if p._p is toc_sect_para._p), -1)
    print("Anchor: TOC sectPr paragraph found at para index", idx2)

    # ── Run fixes ─────────────────────────────────────────────────────────
    add_dedicaces(toc_title_para)
    add_remerciements(toc_title_para)
    update_toc(toc_sect_para)

    doc.save(OUTPUT)
    print("\nSaved -> %s" % OUTPUT)


main()
