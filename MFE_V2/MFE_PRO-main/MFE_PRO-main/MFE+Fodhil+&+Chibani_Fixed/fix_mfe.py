import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph as DocxParagraph

INPUT  = "MFE Fodhil & Chibani.docx"
OUTPUT = "MFE Fodhil & Chibani_Fixed.docx"
BACKUP = "MFE Fodhil & Chibani_BACKUP.docx"
W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def w(tag): return "{%s}%s" % (W_NS, tag)

def run_color(run):
    rpr = run._r.find(w("rPr"))
    if rpr is not None:
        cel = rpr.find(w("color"))
        if cel is not None:
            return cel.get(w("val"), "").upper()
    return ""

def is_red(run): return run_color(run) == "FF0000"

def decolor(run):
    rpr = run._r.find(w("rPr"))
    if rpr is not None:
        cel = rpr.find(w("color"))
        if cel is not None: cel.set(w("val"), "auto")

def para_text(p): return "".join(r.text for r in p.runs if r.text)

def para_full_text(p):
    buf = ""
    for child in p._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            for t in child.findall(w("t")):
                if t.text: buf += t.text
        elif tag == "hyperlink":
            for r in child.findall(w("r")):
                for t in r.findall(w("t")):
                    if t.text: buf += t.text
    return buf

def clear_runs(p):
    for r in p._p.findall(w("r")): p._p.remove(r)

def clear_all(p):
    for tag in (w("r"), w("hyperlink")):
        for el in p._p.findall(tag): p._p.remove(el)

def set_text(p, text, bold=False, italic=False, center=False):
    clear_runs(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def all_paras(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs

def delete_para(p):
    parent = p._p.getparent()
    if parent is not None: parent.remove(p._p)

def ins(ref, text="", bold=False, italic=False, center=False, pb=False):
    new_p = OxmlElement("w:p")
    ref._p.addprevious(new_p)
    para = DocxParagraph(new_p, ref._p.getparent())
    if pb:
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(w("type"), "page")
        run._r.append(br)
    elif text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if center: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


# ── Fix 1 ─────────────────────────────────────────────────────────────────────
def fix_labels(doc):
    print("\n[Fix 1] Table/Figure labels")
    for p in all_paras(doc):
        txt = para_text(p)
        if txt.strip() in ("Table ?????", "Table?????") and "OCPP" not in txt:
            set_text(p, "Table 1.1 – Summary of ISO 15118 Implementation Parameters",
                     bold=True, italic=True, center=True)
            print("  + Table 1.1 fixed")
        elif "Figure ????" in txt:
            set_text(p, "Figure 2.1 – Use Case Diagram: Authentication via Start Button",
                     italic=True, center=True)
            print("  + Figure 2.1 fixed")
        elif "Table ?????" in txt and txt.strip() != "Table ?????":
            set_text(p, "Table 2.1 – Technical Comparison: OCPP 1.6 vs. OCPP 2.0",
                     bold=True, italic=True, center=True)
            print("  + Table 2.1 fixed")


# ── Fix 3a ────────────────────────────────────────────────────────────────────
def fix_vtg(doc):
    print("\n[Fix 3a] Vehicle-to-Grid")
    for p in all_paras(doc):
        runs = p.runs
        for i, run in enumerate(runs):
            if run.text and "Vehicle-to" in run.text and is_red(run):
                if i + 1 < len(runs) and runs[i+1].text and runs[i+1].text.startswith("Grid"):
                    run.text = "Vehicle-to-"
                    decolor(run)
                    print("  + done (split runs)")
                    return
                if "toGrid" in run.text:
                    run.text = run.text.replace("toGrid", "to-Grid")
                    decolor(run)
                    print("  + done (same run)")
                    return


# ── Fix 3b ────────────────────────────────────────────────────────────────────
def fix_tls(doc):
    print("\n[Fix 3b] TLS definition")
    for p in all_paras(doc):
        for run in p.runs:
            if run.text and "TLS-based" in run.text and is_red(run):
                run.text = run.text.replace("TLS-based",
                    "Transport Layer Security (TLS)-based")
                decolor(run)
                print("  + done")
                return


# ── Fix 4 ─────────────────────────────────────────────────────────────────────
def fix_cdr(doc):
    print("\n[Fix 4] CurrentDemandRes citation")
    for p in all_paras(doc):
        txt = para_text(p)
        if "CurrentDemandRes" in txt and "demand response" in txt:
            set_text(p,
                "One of the primary use cases for ISO 15118 in our platform is demand "
                "response driven by photovoltaic (PV) generation. When PV output "
                "fluctuates, the EVSEpi management code adjusts the EVSE current and "
                "voltage limits accordingly. These updated limits are communicated to "
                "the EV via the CurrentDemandRes message, as specified in ISO 15118-2 "
                "[ISO 15118-2, 2022], causing the vehicle to adapt its charge rate "
                "within the subsequent 5-second communication cycle."
            )
            print("  + Rewritten with citation [ISO 15118-2, 2022]")
            return


# ── Fix 2 + 8 ─────────────────────────────────────────────────────────────────
def fix_ocpp(doc):
    print("\n[Fix 2+8] Why OCPP 1.6 rewrite")

    for p in doc.paragraphs:
        if "The choice of OCPP 1.6 over OCPP 2.0 in the Algerian context" in para_text(p):
            set_text(p,
                "Algeria is currently in an early phase of electric vehicle deployment, "
                "characterized by limited charging infrastructure and the absence of an "
                "established smart grid or Vehicle-to-Grid (V2G) ecosystem. Consequently, "
                "the advanced functionalities introduced by OCPP 2.0 — including "
                "dynamic smart charging, bidirectional energy management, and native "
                "ISO 15118 integration — are not operationally required in the "
                "present Algerian context."
            )
            print("  + Intro rewritten")
            break

    dead_bullets = [
        "Algeria is still in early EV deployment stage",
        "No large smart grid integration yet",
        "No V2G infrastructure",
        "Therefore, advanced features of OCPP 2.0 are not necessary",
    ]
    for frag in dead_bullets:
        for p in list(doc.paragraphs):
            if frag in para_text(p):
                delete_para(p)
                print("  + Deleted: %s" % frag[:50])
                break

    for p in doc.paragraphs:
        if "Most low-cost chargers available in developing markets support only OCPP 1.6" in para_text(p):
            set_text(p,
                "In developing markets such as Algeria, the majority of commercially "
                "available low-cost charging hardware supports only OCPP 1.6. "
                "OCPP 2.0-certified equipment remains scarce and significantly more "
                "expensive, making OCPP 1.6 the economically viable and practically "
                "accessible choice for deployment in such contexts."
            )
            print("  + Hardware availability rewritten")
            break
    for frag in ["OCPP 2.0 chargers are more expensive and rare", "OCPP 1.6 is economically suitable"]:
        for p in list(doc.paragraphs):
            if frag in para_text(p): delete_para(p); print("  + Deleted: %s" % frag[:50]); break

    for p in doc.paragraphs:
        if "OCPP 1.6 easier to implement" in para_text(p):
            set_text(p,
                "From a development standpoint, OCPP 1.6 offers considerably greater "
                "simplicity and benefits from an extensive ecosystem of documentation, "
                "tutorials, and open-source tooling. Numerous reference implementations "
                "and protocol simulators are available for OCPP 1.6, facilitating rapid "
                "prototyping and academic experimentation within the scope of this project."
            )
            print("  + Simplicity section rewritten")
            break
    for frag in ["More documentation available", "Many open-source simulators use OCPP 1.6"]:
        for p in list(doc.paragraphs):
            if frag in para_text(p): delete_para(p); print("  + Deleted: %s" % frag[:50]); break

    for p in doc.paragraphs:
        if "OCPP 1.6 works well with low bandwidth" in para_text(p):
            set_text(p,
                "OCPP 1.6 imposes significantly lower communication overhead than "
                "OCPP 2.0, making it more resilient in environments characterized by "
                "limited or intermittent network connectivity. Given Algeria’s "
                "uneven internet infrastructure — particularly outside major urban "
                "centers — OCPP 1.6 provides a more reliable communication layer "
                "for deployed charging stations, whereas OCPP 2.0’s advanced "
                "messaging model would introduce unnecessary complexity in such settings."
            )
            print("  + Network constraints rewritten")
            break
    for frag in ["OCPP 2.0 requires continuous advanced communication",
                 "1.6 is better for unstable internet environments"]:
        for p in list(doc.paragraphs):
            if frag in para_text(p): delete_para(p); print("  + Deleted: %s" % frag[:50]); break

    for p in doc.paragraphs:
        if "Most deployed chargers worldwide still run OCPP 1.6" in para_text(p):
            set_text(p,
                "OCPP 1.6 remains the dominant protocol in deployed charging networks "
                "worldwide. Its widespread adoption ensures interoperability across a "
                "broad range of vendor hardware and backend systems — a critical "
                "factor for any scalable deployment strategy, and particularly relevant "
                "when integrating with existing infrastructure in developing markets."
            )
            print("  + Industry compatibility rewritten")
            break


# ── Fix 5 ─────────────────────────────────────────────────────────────────────
def fix_refs(doc):
    print("\n[Fix 5] Chapter 2 refs -> APA")
    oca = [
        ("Open Charge Alliance. (2019). Open Charge Point Protocol 1.6, Edition 2. "
         "Open Charge Alliance. "
         "https://www.openchargealliance.org/uploads/files/ocpp-1.6.pdf"),
        ("Open Charge Alliance. (2022). Open Charge Point Protocol 2.0.1. "
         "Open Charge Alliance. https://www.openchargealliance.org/"),
    ]
    repl = {
        "luxmanenergy.com": (
            "LuxmanEnergy. (2024). OCPP 1.6 vs OCPP 2.0: A Detailed Comparison for "
            "EV Chargers. Retrieved May 1, 2026, from "
            "https://www.luxmanenergy.com/ocpp-1-6-vs-ocpp-2-0-a-detailed-comparison-for-ev-chargers/"
        ),
        "ampcontrol.io": (
            "AmpControl. (2024). OCPP 1.6 vs OCPP 2.0: A Comprehensive Comparison. "
            "Retrieved May 1, 2026, from "
            "https://www.ampcontrol.io/post/ocpp-1-6-vs-ocpp-2-0-a-comprehensive-comparison"
        ),
    }
    oca_done = False
    for p in doc.paragraphs:
        txt = para_full_text(p)
        for key, apa in repl.items():
            if key in txt:
                if not oca_done:
                    ins(p, oca[1])
                    ins(p, oca[0])
                    oca_done = True
                clear_all(p)
                p.add_run(apa)
                print("  + Replaced %s" % key)
                break


# ── Fix 7 + 8 ─────────────────────────────────────────────────────────────────
def add_lists(doc):
    print("\n[Fix 7+8] Lists + AI declaration")
    ch1 = None
    for p in doc.paragraphs:
        if "CHAPTER 1" in para_text(p) and "ISO 15118" in para_text(p):
            ch1 = p; break
    if not ch1:
        print("  ! Chapter 1 heading not found"); return

    # Insert order reversed: last inserted = appears first before ch1
    # Reading order: List of Figures -> List of Tables -> List of Abbreviations
    #                -> AI Note -> [page break] -> CHAPTER 1

    ins(ch1, pb=True)

    ins(ch1, "")
    ins(ch1,
        "Note on the Use of Artificial Intelligence Tools: In the preparation of this "
        "document, the authors made use of AI-assisted writing tools (including ChatGPT "
        "and GitHub Copilot) to support drafting, language revision, and code assistance. "
        "All content has been reviewed, reformulated, and validated by the authors, who "
        "bear full responsibility for the accuracy, coherence, and originality of the "
        "work presented herein.")
    ins(ch1, "NOTE ON THE USE OF ARTIFICIAL INTELLIGENCE TOOLS", bold=True, center=True)
    ins(ch1, pb=True)

    abbrevs = [
        ("CC",        "Constant Current"),
        ("CSMS",      "Charging Station Management System"),
        ("CV",        "Constant Voltage"),
        ("EIM",       "External Identification Means"),
        ("EV",        "Electric Vehicle"),
        ("EVCC",      "Electric Vehicle Communication Controller"),
        ("EVSE",      "Electric Vehicle Supply Equipment"),
        ("HMI",       "Human-Machine Interface"),
        ("IEC",       "International Electrotechnical Commission"),
        ("IoT",       "Internet of Things"),
        ("ISO",       "International Organization for Standardization"),
        ("OCPP",      "Open Charge Point Protocol"),
        ("PnC",       "Plug & Charge"),
        ("PV",        "Photovoltaic"),
        ("RISE V2G",  "Robust Implementation and Standardized Extensions for V2G"),
        ("SECC",      "Supply Equipment Communication Controller"),
        ("SoC",       "State of Charge"),
        ("TLS",       "Transport Layer Security"),
        ("V2G",       "Vehicle-to-Grid"),
    ]
    ins(ch1, "")
    for abbr, full in reversed(abbrevs):
        ins(ch1, "%-14s%s" % (abbr, full))
    ins(ch1, "")
    ins(ch1, "LIST OF ABBREVIATIONS", bold=True, center=True)
    ins(ch1, pb=True)

    tables_list = [
        ("Table 1.1", "Summary of ISO 15118 Implementation Parameters"),
        ("Table 2.1", "Technical Comparison: OCPP 1.6 vs. OCPP 2.0"),
    ]
    ins(ch1, "")
    for num, title in reversed(tables_list):
        ins(ch1, "%s – %s" % (num, title))
    ins(ch1, "")
    ins(ch1, "LIST OF TABLES", bold=True, center=True)
    ins(ch1, pb=True)

    ins(ch1, "Figure 2.1 – Use Case Diagram: Authentication via Start Button")
    ins(ch1, "")
    ins(ch1, "LIST OF FIGURES", bold=True, center=True)

    print("  + List of Figures inserted")
    print("  + List of Tables inserted")
    print("  + List of Abbreviations inserted (19 entries)")
    print("  + AI tool declaration inserted")


# ── Cleanup ───────────────────────────────────────────────────────────────────
def cleanup(doc):
    print("\n[Cleanup] Residual red formatting")
    n = 0
    for p in all_paras(doc):
        for run in p.runs:
            if is_red(run): decolor(run); n += 1
    print("  + %d red run(s) cleared" % n)


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("Backing up -> %s" % BACKUP)
    shutil.copy2(INPUT, BACKUP)
    doc = Document(INPUT)
    fix_labels(doc)
    fix_vtg(doc)
    fix_tls(doc)
    fix_cdr(doc)
    fix_ocpp(doc)
    fix_refs(doc)
    add_lists(doc)
    cleanup(doc)
    doc.save(OUTPUT)
    print("\nDone -> %s" % OUTPUT)

main()
